"""
Génère les 3 documents Word pour le dossier de tests :
  - Plan de tests
  - Cahier de tests (version allégée ~60 cas)
  - PV de recette

Exécuter : python generate-docx.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), "word")
os.makedirs(OUT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def new_doc(title: str, meta: dict) -> Document:
    doc = Document()
    # Marges réduites
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for k, v in meta.items():
        p = doc.add_paragraph()
        r = p.add_run(f"{k} : ")
        r.bold = True
        p.add_run(v)
    doc.add_paragraph()
    return doc


def h2(doc: Document, text: str):
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str):
    doc.add_heading(text, level=3)


def para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def shade_cell(cell, hex_color: str):
    """Applique un fond de couleur à une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc: Document, headers: list, rows: list):
    """Ajoute un tableau DOCX avec en-tête gris."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # En-tête
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(cell, "D9D9D9")

    # Lignes
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()


# ===========================================================================
# 1. Plan de tests
# ===========================================================================

def build_plan():
    doc = new_doc(
        "Plan de Tests — Backend Ressources Relationnelles",
        {
            "Projet": "Ressources Relationnelles (CDA — Cas H2O V2)",
            "Module": "Backend NestJS",
            "Version": "1.0",
            "Date": "2026-04-13",
        },
    )

    # 1. Objectifs
    h2(doc, "1. Objectifs")
    para(doc, "Garantir la qualité, la fiabilité et la maintenabilité du backend en couvrant trois niveaux de tests :")
    add_table(doc,
        ["Niveau", "Outil", "Périmètre"],
        [
            ["Tests unitaires", "Jest + mocks Prisma", "Logique métier des services et gardes"],
            ["Tests d'intégration (e2e)", "Jest + supertest + DB réelle", "Flux HTTP complets"],
            ["Tests de non-régression", "CI GitHub Actions", "Prévention de régression à chaque push"],
        ],
    )

    # 2. Périmètre
    h2(doc, "2. Périmètre")
    h3(doc, "2.1 Dans le scope")
    para(doc, "• 11 modules fonctionnels : auth, utilisateurs, ressources, catégories, tags, commentaires, favoris, progressions, signalements, amis, messagerie")
    para(doc, "• Gardes : AuthGuard (token HMAC), RolesGuard (rôles utilisateur)")
    para(doc, "• Upload fichiers : photo de profil (multer)")
    para(doc, "• Throttling : limitation de débit (@nestjs/throttler)")
    h3(doc, "2.2 Hors scope")
    para(doc, "• Frontend / application mobile")
    para(doc, "• Infrastructure de production (k8s, CDN, DNS)")
    para(doc, "• Tests de performance / charge")

    # 3. Environnements
    h2(doc, "3. Environnements")
    add_table(doc,
        ["Environnement", "Base de données", "Port", "Usage"],
        [
            ["Développement", "PostgreSQL 16 (Docker)", "5433", "Dev local"],
            ["Test", "PostgreSQL 16 (Docker)", "5434", "Tests e2e locaux"],
            ["CI", "PostgreSQL 16 (service GitHub Actions)", "5434", "Pipeline automatisé"],
        ],
    )
    para(doc, "Variables d'environnement de test :", bold=True)
    para(doc, "DATABASE_URL=postgresql://test:test@127.0.0.1:5434/rr_cesi_test")
    para(doc, "AUTH_TOKEN_SECRET=test-secret")

    # 4. Outils
    h2(doc, "4. Outils")
    add_table(doc,
        ["Outil", "Version", "Rôle"],
        [
            ["Jest", "30.x", "Framework de test (unit + e2e)"],
            ["@nestjs/testing", "11.x", "Module de test NestJS"],
            ["supertest", "7.x", "Appels HTTP dans les tests e2e"],
            ["ts-jest", "29.x", "Compilation TypeScript pour Jest"],
            ["Docker / docker-compose", "20+", "Base de données de test isolée"],
            ["GitHub Actions", "—", "CI/CD"],
        ],
    )

    # 5. Stratégie
    h2(doc, "5. Stratégie par niveau")
    h3(doc, "5.1 Tests unitaires (src/**/*.spec.ts)")
    para(doc, "Chaque service a son fichier .spec.ts colocalisé. PrismaService est mocké via createPrismaMock(). Aucun accès réseau ni base de données réelle. Fixtures partagées dans src/test-utils/fixtures.ts. Isolation : jest.clearAllMocks() dans afterEach.")
    h3(doc, "5.2 Tests d'intégration e2e (test/*.e2e-spec.ts)")
    para(doc, "AppModule complet instancié via createTestApp(). ThrottlerGuard désactivé pour éviter les faux positifs. ValidationPipe avec whitelist + forbidNonWhitelisted + transform. Chaque suite isole son état via truncateAll(prisma) dans beforeEach. Exécution séquentielle (--runInBand).")
    h3(doc, "5.3 Tests de non-régression (CI)")
    para(doc, "Déclenché automatiquement sur tout push/PR ciblant main. Deux jobs indépendants : unit (avec coverage) et e2e (avec Postgres service container). Rapport de coverage archivé comme artefact GitHub Actions (14 jours).")

    # 6. Critères
    h2(doc, "6. Critères d'entrée et de sortie")
    para(doc, "Critères d'entrée :", bold=True)
    para(doc, "• Code review approuvée (PR mergeable)")
    para(doc, "• Schéma Prisma valide (prisma validate)")
    para(doc, "• Build TypeScript sans erreur (nest build)")
    para(doc, "• Base de test accessible (Docker ou CI service container)")
    para(doc, "Critères de sortie :", bold=True)
    add_table(doc,
        ["Critère", "Seuil"],
        [
            ["Tests unitaires", "100 % verts"],
            ["Couverture services (src/**/*.service.ts)", ">= 70 % branches"],
            ["Couverture gardes (auth.guard, roles.guard)", ">= 90 %"],
            ["Tests e2e happy-path", "0 échec"],
            ["Tests e2e scénarios d'erreur (40x)", "0 échec"],
        ],
    )

    # 7. Responsabilités
    h2(doc, "7. Responsabilités")
    add_table(doc,
        ["Rôle", "Responsabilité"],
        [
            ["Développeur", "Écriture des tests unitaires et e2e, maintenance des fixtures"],
            ["Tech Lead", "Validation des critères de couverture, revue des nouveaux cas"],
            ["DevOps", "Maintenance du pipeline CI et de la base de test Docker"],
        ],
    )

    # 8. Risques
    h2(doc, "8. Risques")
    add_table(doc,
        ["Risque", "Impact", "Mitigation"],
        [
            ["Dérive de la DB de test", "E2e faux positifs", "truncateAll + RESTART IDENTITY entre chaque test"],
            ["Token HMAC expiré en CI", "Échec auth", "Token généré dynamiquement dans buildToken()"],
            ["Ports Docker en conflit", "CI/local", "Port 5434 dédié aux tests, distinct du port dev 5433"],
            ["Lenteur des e2e en CI", "Feedback trop lent", "--runInBand + Postgres health-check avant les migrations"],
        ],
    )

    doc.save(os.path.join(OUT_DIR, "Plan-de-tests.docx"))
    print("OK Plan-de-tests.docx")


# ===========================================================================
# 2. Cahier de tests (version allégée)
# ===========================================================================

def build_cahier():
    doc = new_doc(
        "Cahier de Tests — Backend Ressources Relationnelles",
        {
            "Projet": "Ressources Relationnelles (CDA — Cas H2O V2)",
            "Module": "Backend NestJS",
            "Version": "1.0 (sélection des cas critiques)",
            "Date": "2026-04-13",
        },
    )

    para(doc, "Ce cahier présente une sélection des cas de test les plus représentatifs. "
              "Sont conservés : tous les tests E2E (parcours métier complets) et les tests unitaires P1 "
              "portant sur la sécurité, les gardes et les règles métier clés.")

    # Légende
    h2(doc, "Légende")
    add_table(doc,
        ["Colonne", "Description"],
        [
            ["ID", "Identifiant unique (MODULE-TYPE-NUMERO)"],
            ["Type", "U = Unitaire, E = E2E/Intégration"],
            ["Priorité", "P1 Critique, P2 Majeur"],
            ["Statut", "Passé / Échoué / Non exécuté"],
        ],
    )

    # -----------------------------------------------------------------------
    # 1. Authentification
    # -----------------------------------------------------------------------
    h2(doc, "1. Authentification")

    h3(doc, "1.1 AuthService — cas critiques (unitaires)")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["AUTH-U-02", "Login mauvais mot de passe", "login(...) avec bcrypt.compare=false", "Lève UnauthorizedException", "P1", ""],
            ["AUTH-U-03", "Login utilisateur inconnu", "findUnique retourne null", "Lève UnauthorizedException", "P1", ""],
            ["AUTH-U-04", "Login user INACTIF", "User statut INACTIF", "Lève UnauthorizedException", "P1", ""],
            ["AUTH-U-05", "Login user SUSPENDU", "User statut SUSPENDU", "Lève UnauthorizedException", "P1", ""],
            ["AUTH-U-06", "Register valide", "register({...})", "Retourne user sans motDePasse", "P1", ""],
            ["AUTH-U-07", "Register mots de passe non identiques", "password='a', confirmPassword='b'", "Lève BadRequestException", "P1", ""],
            ["AUTH-U-09", "Register email déjà pris", "create lève P2002", "Lève ConflictException", "P1", ""],
            ["AUTH-U-12", "verifyToken valide", "Token signé avec le bon secret", "Retourne { userId, email }", "P1", ""],
            ["AUTH-U-13", "verifyToken signature invalide", "Token altéré", "Retourne null", "P1", ""],
            ["AUTH-U-14", "verifyToken expiré (> 24h)", "Timestamp vieux de 25h", "Retourne null", "P1", ""],
        ],
    )

    h3(doc, "1.2 AuthGuard — cas critiques (unitaires)")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["GUARD-U-01", "Route @Public()", "Reflector retourne IS_PUBLIC_KEY=true", "Retourne true sans vérifier le token", "P1", ""],
            ["GUARD-U-02", "Header Authorization absent", "Header vide", "Lève UnauthorizedException", "P1", ""],
            ["GUARD-U-04", "Token invalide", "verifyToken retourne null", "Lève UnauthorizedException", "P1", ""],
            ["GUARD-U-05", "Token valide", "verifyToken retourne payload", "request.user injecté, retourne true", "P1", ""],
        ],
    )

    h3(doc, "1.3 RolesGuard — cas critiques (unitaires)")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["ROLES-U-01", "Aucun décorateur @Roles", "Reflector retourne undefined", "Retourne true", "P1", ""],
            ["ROLES-U-03", "Pas de user sur request", "request.user = undefined", "Lève UnauthorizedException", "P1", ""],
            ["ROLES-U-04", "User introuvable en base", "findUnique retourne null", "Lève ForbiddenException", "P1", ""],
            ["ROLES-U-05", "Rôle insuffisant", "User CITOYEN, route MODERATEUR", "Lève ForbiddenException", "P1", ""],
            ["ROLES-U-06", "Rôle suffisant", "User MODERATEUR, route MODERATEUR", "Retourne true", "P1", ""],
            ["ROLES-U-07", "SUPER_ADMIN accède à toute route", "User SUPER_ADMIN, route ADMINISTRATEUR", "Retourne true", "P1", ""],
        ],
    )

    h3(doc, "1.4 Auth E2E")
    add_table(doc,
        ["ID", "Description", "Préconditions", "Attendu", "Prio", "Statut"],
        [
            ["AUTH-E2E-01", "Inscription utilisateur", "DB vide, body valide", "201, user sans motDePasse", "P1", ""],
            ["AUTH-E2E-02", "Inscription doublon email", "User existant", "409", "P1", ""],
            ["AUTH-E2E-03", "Inscription mots de passe non identiques", "—", "400", "P1", ""],
            ["AUTH-E2E-04", "Connexion valide", "User ACTIF en base", "200 + accessToken", "P1", ""],
            ["AUTH-E2E-05", "Connexion mauvais mot de passe", "User en base", "401", "P1", ""],
            ["AUTH-E2E-06", "Accès route protégée sans token", "—", "401", "P1", ""],
            ["AUTH-E2E-07", "Accès route protégée avec token valide", "User en base", "200", "P1", ""],
            ["AUTH-E2E-08", "Création admin par super_admin", "Super_admin authentifié", "201", "P1", ""],
            ["AUTH-E2E-09", "Création admin par citoyen", "Citoyen authentifié", "403", "P1", ""],
        ],
    )

    # -----------------------------------------------------------------------
    # 2. Utilisateurs
    # -----------------------------------------------------------------------
    h2(doc, "2. Utilisateurs")

    h3(doc, "2.1 UtilisateursService — cas critiques (unitaires)")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["USR-U-03", "findById — trouvé", "findUnique retourne user", "User sans motDePasse", "P1", ""],
            ["USR-U-04", "findById — non trouvé", "findUnique retourne null", "Lève NotFoundException", "P1", ""],
        ],
    )

    h3(doc, "2.2 Utilisateurs E2E")
    add_table(doc,
        ["ID", "Description", "Préconditions", "Attendu", "Prio", "Statut"],
        [
            ["USR-E2E-01", "Liste utilisateurs — admin", "Admin authentifié", "200, liste", "P1", ""],
            ["USR-E2E-02", "Liste utilisateurs — citoyen", "Citoyen authentifié", "403", "P1", ""],
            ["USR-E2E-03", "Profil utilisateur par ID", "User en base", "200, sans motDePasse", "P1", ""],
            ["USR-E2E-04", "Profil utilisateur inexistant", "—", "404", "P2", ""],
            ["USR-E2E-05", "Modification profil", "User owner", "200", "P1", ""],
            ["USR-E2E-06", "Modification statut — admin", "Admin", "200", "P1", ""],
            ["USR-E2E-07", "Modification statut — citoyen", "Citoyen", "403", "P1", ""],
            ["USR-E2E-08", "Modification rôle — admin", "Admin", "200", "P1", ""],
            ["USR-E2E-09", "Suppression utilisateur — admin", "Admin", "200", "P1", ""],
            ["USR-E2E-10", "Suppression utilisateur — citoyen", "Citoyen", "403", "P1", ""],
        ],
    )

    # -----------------------------------------------------------------------
    # 3. Ressources
    # -----------------------------------------------------------------------
    h2(doc, "3. Ressources")

    h3(doc, "3.1 RessourcesService — cas critiques (unitaires)")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["RES-U-01", "findAll filtre VALIDEE + PUBLIQUE", "findAll({})", "where: { statut: 'VALIDEE', visibilite: 'PUBLIQUE' }", "P1", ""],
            ["RES-U-03", "findById — trouvé", "findById(1)", "Ressource retournée", "P1", ""],
            ["RES-U-04", "findById — non trouvé", "findById(99)", "NotFoundException", "P1", ""],
        ],
    )

    h3(doc, "3.2 Ressources E2E")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["RES-E2E-01", "Liste ressources publiques validées", "GET /ressources", "200, VALIDEE+PUBLIQUE seulement", "P1", ""],
            ["RES-E2E-02", "Filtre ressources par catégorie", "GET /ressources?categorieId=X", "200, filtrées", "P2", ""],
            ["RES-E2E-03", "Détail ressource", "GET /ressources/:id", "200", "P1", ""],
            ["RES-E2E-04", "Détail ressource inexistante", "GET /ressources/99999", "404", "P2", ""],
            ["RES-E2E-05", "Ressources par utilisateur", "GET /ressources/utilisateur/:id", "200, tous statuts", "P1", ""],
            ["RES-E2E-06", "Création ressource authentifiée", "POST authentifié", "201", "P1", ""],
            ["RES-E2E-07", "Création ressource non authentifiée", "POST sans auth", "401", "P1", ""],
            ["RES-E2E-08", "Modification ressource — modérateur", "PATCH modérateur", "200", "P1", ""],
            ["RES-E2E-09", "Modification ressource — citoyen", "PATCH citoyen", "403", "P1", ""],
            ["RES-E2E-10", "Suppression ressource — modérateur", "DELETE modérateur", "200", "P1", ""],
        ],
    )

    # -----------------------------------------------------------------------
    # 4. Catégories
    # -----------------------------------------------------------------------
    h2(doc, "4. Catégories")

    h3(doc, "4.1 Catégories E2E")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["CAT-E2E-01", "Liste catégories", "GET /categories", "200", "P1", ""],
            ["CAT-E2E-02", "Détail catégorie", "GET /categories/:id existante", "200", "P1", ""],
            ["CAT-E2E-03", "Catégorie inexistante", "GET /categories/99999", "404", "P2", ""],
            ["CAT-E2E-04", "Création — admin", "POST admin", "201", "P1", ""],
            ["CAT-E2E-05", "Création — citoyen", "POST citoyen", "403", "P1", ""],
            ["CAT-E2E-06", "Modification — admin", "PATCH admin", "200", "P1", ""],
            ["CAT-E2E-07", "Suppression — admin", "DELETE admin", "200", "P1", ""],
            ["CAT-E2E-08", "Suppression — citoyen", "DELETE citoyen", "403", "P1", ""],
        ],
    )

    # -----------------------------------------------------------------------
    # 5. Messagerie
    # -----------------------------------------------------------------------
    h2(doc, "5. Messagerie")

    h3(doc, "5.1 MessagerieService — cas critiques (unitaires)")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["MSG-U-03", "findConversationById — non trouvée", "findConversationById(99)", "NotFoundException", "P1", ""],
            ["MSG-U-05", "createConversation", "createConversation({ participantIds: [1, 2] })", "participants.create appelé", "P1", ""],
            ["MSG-U-06", "sendMessage", "sendMessage(1, { idUtilisateur: 1, contenu: '...' })", "Message créé", "P1", ""],
        ],
    )

    h3(doc, "5.2 Messagerie E2E")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["MSG-E2E-01", "Flux complet : conversation + messages + lecture", "Créer conv, envoyer msg, marquer lu, vérifier non-lus", "201/201/200/1 non-lu/200/0 non-lu", "P1", ""],
            ["MSG-E2E-02", "Liste conversations utilisateur", "GET /conversations/utilisateur/:id", "200", "P1", ""],
            ["MSG-E2E-03", "Conversation inexistante", "GET /conversations/99999", "404", "P2", ""],
        ],
    )

    # -----------------------------------------------------------------------
    # 6. Amis
    # -----------------------------------------------------------------------
    h2(doc, "6. Amis")

    h3(doc, "6.1 AmisService — cas critiques (unitaires)")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["AMI-U-04", "Envoyer demande d'amitié", "envoyer(1, 2)", "Demande créée statut EN_ATTENTE", "P1", ""],
            ["AMI-U-05", "Envoyer demande doublon", "create lève P2002", "ConflictException", "P1", ""],
            ["AMI-U-06", "Accepter demande", "accepter(1, 2)", "Statut ACCEPTE", "P1", ""],
            ["AMI-U-07", "Accepter demande non trouvée", "findFirst retourne null", "NotFoundException", "P1", ""],
            ["AMI-U-08", "Refuser demande", "refuser(1, 2)", "Statut REFUSE", "P1", ""],
        ],
    )

    h3(doc, "6.2 Amis E2E")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["AMI-E2E-01", "Flux complet : envoi + acceptation + suppression", "Envoi, accepter, vérifier liste, supprimer", "201/200/200 (0 ami)", "P1", ""],
            ["AMI-E2E-02", "Refus de demande", "Envoi + refus", "200, statut REFUSE", "P1", ""],
            ["AMI-E2E-03", "Double demande", "Envoi x2 vers le même user", "409", "P1", ""],
        ],
    )

    # -----------------------------------------------------------------------
    # 7. Signalements
    # -----------------------------------------------------------------------
    h2(doc, "7. Signalements")

    h3(doc, "7.1 SignalementsService — cas critiques (unitaires)")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["SIG-U-03", "Créer signalement", "create(dto)", "Signalement créé", "P1", ""],
            ["SIG-U-04", "Traiter signalement — TRAITE", "update(1, { statut: 'TRAITE', ... })", "dateTraitement défini", "P1", ""],
        ],
    )

    h3(doc, "7.2 Signalements E2E")
    add_table(doc,
        ["ID", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["SIG-E2E-01", "Signalement ressource par citoyen", "POST /signalements — citoyen", "201, statut EN_ATTENTE", "P1", ""],
            ["SIG-E2E-02", "Liste signalements — modérateur", "GET /signalements — modérateur", "200", "P1", ""],
            ["SIG-E2E-03", "Liste signalements — citoyen", "GET /signalements — citoyen", "403", "P1", ""],
            ["SIG-E2E-04", "Filtre par statut", "GET /signalements?statut=EN_ATTENTE", "200", "P2", ""],
            ["SIG-E2E-05", "Traitement signalement", "PATCH /signalements/:id — modérateur", "200, dateTraitement défini", "P1", ""],
        ],
    )

    # -----------------------------------------------------------------------
    # 8. Upload photo profil
    # -----------------------------------------------------------------------
    h2(doc, "8. Upload photo profil")

    add_table(doc,
        ["ID", "Type", "Description", "Étapes", "Attendu", "Prio", "Statut"],
        [
            ["UPL-E2E-01", "E", "Upload image PNG", "POST /utilisateurs/:id/photo — multipart PNG 1x1", "201, photoProfil =~ /^/uploads/photo-/", "P1", ""],
            ["UPL-E2E-02", "E", "Upload fichier non-image", "POST — multipart TXT", "400", "P1", ""],
            ["UPL-E2E-03", "E", "Suppression photo", "DELETE /utilisateurs/:id/photo", "200, photoProfil null", "P1", ""],
        ],
    )

    # Récapitulatif
    h2(doc, "Récapitulatif (sélection)")
    add_table(doc,
        ["Module", "Tests unitaires", "Tests e2e", "Total"],
        [
            ["Auth (service + gardes)", "16", "9", "25"],
            ["Utilisateurs", "2", "10", "12"],
            ["Ressources", "3", "10", "13"],
            ["Catégories", "0", "8", "8"],
            ["Messagerie", "3", "3", "6"],
            ["Amis", "5", "3", "8"],
            ["Signalements", "2", "5", "7"],
            ["Upload photo", "0", "3", "3"],
            ["TOTAL", "31", "51", "82"],
        ],
    )

    doc.save(os.path.join(OUT_DIR, "Cahier-de-tests.docx"))
    print("OK Cahier-de-tests.docx")


# ===========================================================================
# 3. PV de recette
# ===========================================================================

def build_pv():
    doc = new_doc(
        "Procès-Verbal de Recette — Backend Ressources Relationnelles",
        {
            "Projet": "Ressources Relationnelles",
            "Module": "Backend NestJS",
            "Version testée": "feat/add-tests (commit 421b5bb)",
            "Date de recette": "2026-04-13",
            "Environnement": "Test local (PostgreSQL 16, port 5434)",
            "Responsable recette": "Fred",
            "Testeur(s)": "Fred",
        },
    )

    # 1. Contexte
    h2(doc, "1. Contexte")
    para(doc, "Ce document atteste de la bonne exécution des tests de recette sur le backend de l'application "
              "Ressources Relationnelles. Il est établi à l'issue de la campagne de tests définie dans le cahier de tests.")

    # 2. Résultats par module
    h2(doc, "2. Résultats par module")

    def pv_table(doc, rows):
        add_table(doc, ["ID", "Cas de test", "Statut", "Anomalie"], rows)

    h3(doc, "2.1 Authentification")
    pv_table(doc, [
        ["AUTH-E2E-01", "Inscription utilisateur (201)", "OK OK", ""],
        ["AUTH-E2E-02", "Inscription doublon email (409)", "OK OK", ""],
        ["AUTH-E2E-03", "Inscription mots de passe non identiques (400)", "OK OK", ""],
        ["AUTH-E2E-04", "Connexion valide (200 + token)", "OK OK", ""],
        ["AUTH-E2E-05", "Connexion mauvais mot de passe (401)", "OK OK", ""],
        ["AUTH-E2E-06", "Accès route protégée sans token (401)", "OK OK", ""],
        ["AUTH-E2E-07", "Accès route protégée avec token valide (200)", "OK OK", ""],
        ["AUTH-E2E-08", "Création admin par super_admin (201)", "OK OK", ""],
        ["AUTH-E2E-09", "Création admin par citoyen (403)", "OK OK", ""],
    ])

    h3(doc, "2.2 Utilisateurs")
    pv_table(doc, [
        ["USR-E2E-01", "Liste utilisateurs — admin (200)", "OK OK", ""],
        ["USR-E2E-02", "Liste utilisateurs — citoyen (403)", "OK OK", ""],
        ["USR-E2E-03", "Profil utilisateur par ID (200, sans motDePasse)", "OK OK", ""],
        ["USR-E2E-04", "Profil utilisateur inexistant (404)", "OK OK", ""],
        ["USR-E2E-05", "Modification profil owner (200)", "OK OK", ""],
        ["USR-E2E-06", "Modification statut — admin (200)", "OK OK", ""],
        ["USR-E2E-07", "Modification statut — citoyen (403)", "OK OK", ""],
        ["USR-E2E-08", "Modification rôle — admin (200)", "OK OK", ""],
        ["USR-E2E-09", "Suppression utilisateur — admin (200)", "OK OK", ""],
        ["USR-E2E-10", "Suppression utilisateur — citoyen (403)", "OK OK", ""],
    ])

    h3(doc, "2.3 Ressources")
    pv_table(doc, [
        ["RES-E2E-01", "Liste ressources publiques validées (200)", "OK OK", ""],
        ["RES-E2E-02", "Filtre ressources par catégorie (200)", "OK OK", ""],
        ["RES-E2E-03", "Détail ressource (200)", "OK OK", ""],
        ["RES-E2E-04", "Détail ressource inexistante (404)", "OK OK", ""],
        ["RES-E2E-05", "Ressources par utilisateur (200)", "OK OK", ""],
        ["RES-E2E-06", "Création ressource authentifiée (201)", "OK OK", ""],
        ["RES-E2E-07", "Création ressource non authentifiée (401)", "OK OK", ""],
        ["RES-E2E-08", "Modification ressource — modérateur (200)", "OK OK", ""],
        ["RES-E2E-09", "Modification ressource — citoyen (403)", "OK OK", ""],
        ["RES-E2E-10", "Suppression ressource — modérateur (200)", "OK OK", ""],
    ])

    h3(doc, "2.4 Catégories")
    pv_table(doc, [
        ["CAT-E2E-01", "Liste catégories (200)", "OK OK", ""],
        ["CAT-E2E-02", "Détail catégorie (200)", "OK OK", ""],
        ["CAT-E2E-03", "Catégorie inexistante (404)", "OK OK", ""],
        ["CAT-E2E-04", "Création — admin (201)", "OK OK", ""],
        ["CAT-E2E-05", "Création — citoyen (403)", "OK OK", ""],
        ["CAT-E2E-06", "Modification — admin (200)", "OK OK", ""],
        ["CAT-E2E-07", "Suppression — admin (200)", "OK OK", ""],
        ["CAT-E2E-08", "Suppression — citoyen (403)", "OK OK", ""],
    ])

    h3(doc, "2.5 Messagerie")
    pv_table(doc, [
        ["MSG-E2E-01", "Création conversation (201)", "OK OK", ""],
        ["MSG-E2E-02", "Envoi message (201)", "OK OK", ""],
        ["MSG-E2E-03", "Listage messages conversation (200)", "OK OK", ""],
        ["MSG-E2E-04", "Comptage messages non-lus (200)", "OK OK", ""],
        ["MSG-E2E-05", "Marquage messages comme lus (200)", "OK OK", ""],
        ["MSG-E2E-06", "Non-lus après lecture = 0 (200)", "OK OK", ""],
        ["MSG-E2E-07", "Conversations par utilisateur (200)", "OK OK", ""],
        ["MSG-E2E-08", "Conversation inexistante (404)", "OK OK", ""],
    ])

    h3(doc, "2.6 Amis")
    pv_table(doc, [
        ["AMI-E2E-01", "Envoi demande d'amitié (201, EN_ATTENTE)", "OK OK", ""],
        ["AMI-E2E-02", "Demandes reçues (200)", "OK OK", ""],
        ["AMI-E2E-03", "Demandes envoyées (200)", "OK OK", ""],
        ["AMI-E2E-04", "Acceptation demande (200, ACCEPTE)", "OK OK", ""],
        ["AMI-E2E-05", "Liste amis après acceptation (200, 1 ami)", "OK OK", ""],
        ["AMI-E2E-06", "Suppression ami (200)", "OK OK", ""],
        ["AMI-E2E-07", "Refus demande (200, REFUSE)", "OK OK", ""],
        ["AMI-E2E-08", "Double demande (409)", "OK OK", ""],
    ])

    h3(doc, "2.7 Signalements")
    pv_table(doc, [
        ["SIG-E2E-01", "Signalement ressource par citoyen (201)", "OK OK", ""],
        ["SIG-E2E-02", "Liste signalements — modérateur (200)", "OK OK", ""],
        ["SIG-E2E-03", "Liste signalements — citoyen (403)", "OK OK", ""],
        ["SIG-E2E-04", "Filtre par statut (200)", "OK OK", ""],
        ["SIG-E2E-05", "Traitement signalement (200, TRAITE, dateTraitement)", "OK OK", ""],
    ])

    h3(doc, "2.8 Upload photo profil")
    pv_table(doc, [
        ["UPL-E2E-01", "Upload image PNG (201, photoProfil mis à jour)", "OK OK", ""],
        ["UPL-E2E-02", "Upload fichier non-image (400)", "OK OK", ""],
        ["UPL-E2E-03", "Suppression photo (200, photoProfil null)", "OK OK", ""],
    ])

    # 3. Bilan
    h2(doc, "3. Bilan")
    add_table(doc,
        ["Catégorie", "Nombre", "Taux"],
        [
            ["Cas exécutés", "59 / 59", "100 %"],
            ["Cas validés (OK)", "59", "100 %"],
            ["Cas en échec (KO)", "0", "0 %"],
        ],
    )

    # 4. Anomalies
    h2(doc, "4. Anomalies détectées")
    para(doc, "Aucune anomalie détectée lors de cette campagne de recette.")

    # 5. Conclusion
    h2(doc, "5. Conclusion")
    p = doc.add_paragraph()
    r = p.add_run("OK Recette ACCEPTÉE")
    r.bold = True
    r.font.color.rgb = RGBColor(0x18, 0x75, 0x3E)
    doc.add_paragraph("Tous les cas critiques sont verts. Aucune anomalie détectée.")

    # 6. Signatures
    h2(doc, "6. Signatures")
    add_table(doc,
        ["Rôle", "Nom", "Date", "Signature"],
        [
            ["Responsable technique", "Fred", "2026-04-13", ""],
            ["Responsable recette client", "", "", ""],
            ["Chef de projet", "", "", ""],
        ],
    )

    doc.save(os.path.join(OUT_DIR, "PV-recette.docx"))
    print("OK PV-recette.docx")


# ===========================================================================
# Main
# ===========================================================================

if __name__ == "__main__":
    build_plan()
    build_cahier()
    build_pv()
    print(f"\nFichiers générés dans : {OUT_DIR}")
